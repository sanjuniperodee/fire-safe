import json
import uuid
from time import sleep
from docx.api import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

from helpers.logger import log_exception
from helpers.utils import convert_datetime

TMP_PATH = "/code/generators/templates"
DOC_PATH = "/code/generators/documents"


def permissionForHotWork(data):
    try:
        json_string = json.dumps(data, default=convert_datetime, indent=2)
        parsed_object = json.loads(json_string)
        data_str = str(datetime.now().date())

        document = Document(f"{TMP_PATH}/РАЗРЕШЕНИЕ на производство огневых работ template.docx")
        file_path = f"{DOC_PATH}/РАЗРЕШЕНИЕ на производство огневых работ {data_str}.docx"

        document.styles.add_style("CommentsStyle", WD_STYLE_TYPE.PARAGRAPH)
        style = document.styles["CommentsStyle"]
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)

        for p in document.paragraphs:
            p.style = document.styles["CommentsStyle"]
            for key, value in parsed_object.items():
                placeholder = f"${{{key}}}"
                if type(value) is dict:
                    for key2, value2 in value.items():
                        placeholder = f"${{{key}.{key2}}}"
                        if placeholder in p.text:
                            p.text = p.text.replace(placeholder, str(value2))
                if "Date" in key:
                    p.text = p.text.replace(placeholder,
                                            str(datetime.fromisoformat(value.replace("Z", "+00:00")).date()))
                if placeholder in p.text:
                    p.text = p.text.replace(placeholder, str(value))

            if "${documentNumber}" in p.text:
                p.text = p.text.replace("${documentNumber}", str(uuid.uuid4()))

            if "${documentCreatedAt}" in p.text:
                p.text = p.text.replace("${documentCreatedAt}", str(datetime.now().date()))

            if "${formattedIssuedMember}" in p.text:
                p.text = p.text.replace(
                    "${formattedIssuedMember}",
                    parsed_object['issuedMember']['name'].split()[1][0] + ". " +
                    parsed_object['issuedMember']['name'].split()[0]
                )

            if "${formattedCarryOutWorker}" in p.text:
                p.text = p.text.replace(
                    "${formattedCarryOutWorker}",
                    parsed_object['carryOutWorker']['name'].split()[1][0] + ". " +
                    parsed_object['carryOutWorker']['name'].split()[0]
                )
            document.save(file_path)

        return file_path
    except Exception as e:
        log_exception(e, 'Error in permissionForHotWork')
        return False


def certificateOfInspectionOfFireAutomaticSystemsAndInstallations(data):
    try:
        json_string = json.dumps(data, default=convert_datetime, indent=2)
        parsed_object = json.loads(json_string)
        data_str = str(datetime.now().date())

        document = Document(f"{TMP_PATH}/АКТ освидетельствования систем и установок пожарной автоматики template.docx")
        file_path = f"{DOC_PATH}/АКТ освидетельствования систем и установок пожарной автоматики {data_str}.docx"

        document.styles.add_style("CommentsStyle", WD_STYLE_TYPE.PARAGRAPH)
        style = document.styles["CommentsStyle"]
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)

        for p in document.paragraphs:
            p.style = document.styles["CommentsStyle"]
            for key, value in parsed_object.items():
                placeholder = f"${{{key}}}"
                if type(value) is dict:
                    for key2, value2 in value.items():
                        placeholder = f"${{{key}.{key2}}}"
                        if placeholder in p.text:
                            p.text = p.text.replace(placeholder, str(value2))
                if "carriedOut" in key:
                    p.text = p.text.replace(
                        placeholder,
                        str(datetime.fromisoformat(value.replace("Z", "+00:00")).date())
                    )
                if placeholder in p.text:
                    p.text = p.text.replace(placeholder, str(value))

            if "${documentNumber}" in p.text:
                p.text = p.text.replace("${documentNumber}", str(uuid.uuid4()))

            if "${documentCreatedAt}" in p.text:
                p.text = p.text.replace("${documentCreatedAt}", str(datetime.now().date()))

            if "${formattedChairmanOfCommision}" in p.text:
                p.text = p.text.replace(
                    "${formattedChairmanOfCommision}",
                    parsed_object['chairmanOfCommision']['name'].split()[1][0] + ". " +
                    parsed_object['chairmanOfCommision']['name'].split()[0]
                )

            if "${commisionMembers}" in p.text:
                resultString = ''
                for member in parsed_object['commissionMembers']:
                    resultString += (member['IIN'] + ', ' + member['name'] + ', ' + member['jobTitle'] + ', ' +
                                     member['BIN'] + ', ' + member['organizationName'] + '\n')
                p.text = p.text.replace("${commisionMembers}", resultString)

            if "${formattedCommisionMembers}" in p.text:
                resultString = ''
                for member in parsed_object['commissionMembers']:
                    resultString += member['name'].split()[1][0] + ". " + member['name'].split()[0] + '\n'
                p.text = p.text.replace("${formattedCommisionMembers}", resultString)

            document.save(file_path)

        return file_path
    except Exception as e:
        log_exception(e, 'Error in certificateOfInspectionOfFireAutomaticSystemsAndInstallations')
        return False


def testingInternalFireFightingWaterSupplySystem(data):
    try:
        json_string = json.dumps(data, default=convert_datetime, indent=2)
        parsed_object = json.loads(json_string)
        data_str = str(datetime.now().date())

        document = Document(
            f"{TMP_PATH}/АКТ ИСПЫТАНИЙ систем внутреннего противопожарного водоснабжения на работоспособность template.docx")
        file_path = f"{DOC_PATH}/АКТ ИСПЫТАНИЙ систем внутреннего противопожарного водоснабжения на работоспособность {data_str}.docx"

        document.styles.add_style("CommentsStyle", WD_STYLE_TYPE.PARAGRAPH)
        style = document.styles["CommentsStyle"]
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)

        for p in document.paragraphs:
            p.style = document.styles["CommentsStyle"]
            for key, value in parsed_object.items():
                placeholder = f"${{{key}}}"
                if type(value) is dict:
                    for key2, value2 in value.items():
                        placeholder = f"${{{key}.{key2}}}"
                        if placeholder in p.text:
                            p.text = p.text.replace(placeholder, str(value2))

                if "testTime" in key:
                    time_placeholder = "${testTime}"
                    if time_placeholder in p.text:
                        p.text = p.text.replace(
                            time_placeholder,
                            str(datetime.fromisoformat(value.replace("Z", "+00:00")).time())
                        )
                elif "testDate" in key:
                    date_placeholder = "${testDate}"
                    if date_placeholder in p.text:
                        p.text = p.text.replace(
                            date_placeholder,
                            str(datetime.fromisoformat(value.replace("Z", "+00:00")).date())
                        )

                if placeholder in p.text:
                    p.text = p.text.replace(placeholder, str(value))

            if "${documentNumber}" in p.text:
                p.text = p.text.replace("${documentNumber}", str(uuid.uuid4()))

            if "${documentCreatedAt}" in p.text:
                p.text = p.text.replace("${documentCreatedAt}", str(datetime.now().date()))

            if "${formattedChairmanOfCommision}" in p.text:
                p.text = p.text.replace(
                    "${formattedChairmanOfCommision}",
                    parsed_object['chairmanOfCommision']['name'].split()[1][0] + ". " +
                    parsed_object['chairmanOfCommision']['name'].split()[0]
                )

            if "${commisionMembers}" in p.text:
                resultString = ''
                for member in parsed_object['commissionMembers']:
                    resultString += (member['IIN'] + ', ' + member['name'] + ', ' + member['jobTitle'] + ', ' +
                                     member['BIN'] + ', ' + member['organizationName'] + '\n')
                p.text = p.text.replace("${commisionMembers}", resultString)

            if "${formattedCommisionMembers}" in p.text:
                resultString = ''
                for member in parsed_object['commissionMembers']:
                    resultString += member['name'].split()[1][0] + ". " + member['name'].split()[0] + '\n'
                p.text = p.text.replace("${formattedCommisionMembers}", resultString)

            document.save(file_path)

        return file_path
    except Exception as e:
        log_exception(e, 'Error in testingInternalFireFightingWaterSupplySystem')
        return False


def actOfInspectionWaterSupplyNetworkForWaterDischarge(data):
    try:
        json_string = json.dumps(data, default=convert_datetime, indent=2)
        parsed_object = json.loads(json_string)
        data_str = str(datetime.now().date())

        document = Document(f"{TMP_PATH}/АКТ ОБСЛЕДОВАНИЯ водопроводной сети на водоотдачу template.docx")
        file_path = f"{DOC_PATH}/АКТ ОБСЛЕДОВАНИЯ водопроводной сети на водоотдачу {data_str}.docx"

        document.styles.add_style("CommentsStyle", WD_STYLE_TYPE.PARAGRAPH)
        style = document.styles["CommentsStyle"]
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)

        for p in document.paragraphs:
            p.style = document.styles["CommentsStyle"]
            for key, value in parsed_object.items():
                placeholder = f"${{{key}}}"
                if type(value) is dict:
                    for key2, value2 in value.items():
                        placeholder = f"${{{key}.{key2}}}"
                        if placeholder in p.text:
                            p.text = p.text.replace(placeholder, str(value2))

                if placeholder in p.text:
                    p.text = p.text.replace(placeholder, str(value))

            if "${documentNumber}" in p.text:
                p.text = p.text.replace("${documentNumber}", str(uuid.uuid4()))

            if "${documentCreatedAt}" in p.text:
                p.text = p.text.replace("${documentCreatedAt}", str(datetime.now().date()))

            if "${formattedChairmanOfCommision}" in p.text:
                p.text = p.text.replace(
                    "${formattedChairmanOfCommision}",
                    parsed_object['chairmanOfCommision']['name'].split()[1][0] + ". " +
                    parsed_object['chairmanOfCommision']['name'].split()[0]
                )

            if "${commisionMembers}" in p.text:
                resultString = ''
                for member in parsed_object['commissionMembers']:
                    resultString += (member['IIN'] + ', ' + member['name'] + ', ' + member['jobTitle'] + ', ' +
                                     member['BIN'] + ', ' + member['organizationName'] + '\n')
                p.text = p.text.replace("${commisionMembers}", resultString)

            if "${formattedCommisionMembers}" in p.text:
                resultString = ''
                for member in parsed_object['commissionMembers']:
                    resultString += member['name'].split()[1][0] + ". " + member['name'].split()[0] + '\n'
                p.text = p.text.replace("${formattedCommisionMembers}", resultString)

            document.save(file_path)

        return file_path
    except Exception as e:
        log_exception(e, 'Error in actOfInspectionWaterSupplyNetworkForWaterDischarge')
        return False


def actOfInspectionFireHydrants(data):
    try:
        json_string = json.dumps(data, default=convert_datetime, indent=2)
        parsed_object = json.loads(json_string)
        data_str = str(datetime.now().date())

        document = Document(f"{TMP_PATH}/АКТ ОБСЛЕДОВАНИЯ пожарных гидрантов template.docx")
        file_path = f"{DOC_PATH}/АКТ ОБСЛЕДОВАНИЯ пожарных гидрантов {data_str}.docx"

        document.styles.add_style("CommentsStyle", WD_STYLE_TYPE.PARAGRAPH)
        style = document.styles["CommentsStyle"]
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)

        for p in document.paragraphs:
            p.style = document.styles["CommentsStyle"]
            for key, value in parsed_object.items():
                placeholder = f"${{{key}}}"
                if type(value) is dict:
                    for key2, value2 in value.items():
                        placeholder = f"${{{key}.{key2}}}"
                        if placeholder in p.text:
                            p.text = p.text.replace(placeholder, str(value2))

                if placeholder in p.text:
                    p.text = p.text.replace(placeholder, str(value))

            if "${documentNumber}" in p.text:
                p.text = p.text.replace("${documentNumber}", str(uuid.uuid4()))

            if "${documentCreatedAt}" in p.text:
                p.text = p.text.replace("${documentCreatedAt}", str(datetime.now().date()))

            if "${formattedChairmanOfCommision}" in p.text:
                p.text = p.text.replace(
                    "${formattedChairmanOfCommision}",
                    parsed_object['chairmanOfCommision']['name'].split()[1][0] + ". " +
                    parsed_object['chairmanOfCommision']['name'].split()[0]
                )

            if "${commisionMembers}" in p.text:
                resultString = ''
                for member in parsed_object['commissionMembers']:
                    resultString += (member['IIN'] + ', ' + member['name'] + ', ' + member['jobTitle'] + ', ' +
                                     member['BIN'] + ', ' + member['organizationName'] + '\n')
                p.text = p.text.replace("${commisionMembers}", resultString)

            if "${formattedCommisionMembers}" in p.text:
                resultString = ''
                for member in parsed_object['commissionMembers']:
                    resultString += member['name'].split()[1][0] + ". " + member['name'].split()[0] + '\n'
                p.text = p.text.replace("${formattedCommisionMembers}", resultString)

            document.save(file_path)

        return file_path
    except Exception as e:
        log_exception(e, 'Error in actOfInspectionFireHydrants')
        return False


def actOfCommissioningFireAutomationSystemsAndInstallations(data):
    try:
        json_string = json.dumps(data, default=convert_datetime, indent=2)
        parsed_object = json.loads(json_string)
        data_str = str(datetime.now().date())

        document = Document(
            f"{TMP_PATH}/АКТ приемки в эксплуатацию систем и установок пожарной автоматики template.docx")
        file_path = f"{DOC_PATH}/АКТ приемки в эксплуатацию систем и установок пожарной автоматики {data_str}.docx"

        document.styles.add_style("CommentsStyle", WD_STYLE_TYPE.PARAGRAPH)
        style = document.styles["CommentsStyle"]
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)

        for p in document.paragraphs:
            p.style = document.styles["CommentsStyle"]
            for key, value in parsed_object.items():
                placeholder = f"${{{key}}}"
                if type(value) is dict:
                    for key2, value2 in value.items():
                        placeholder = f"${{{key}.{key2}}}"
                        if placeholder in p.text:
                            p.text = p.text.replace(placeholder, str(value2))
                if "carriedOut" or "acceptedForOperationFrom" in key:
                    p.text = p.text.replace(
                        placeholder,
                        str(datetime.fromisoformat(value.replace("Z", "+00:00")).date())
                    )
                if placeholder in p.text:
                    p.text = p.text.replace(placeholder, str(value))

            if "${documentNumber}" in p.text:
                p.text = p.text.replace("${documentNumber}", str(uuid.uuid4()))

            if "${documentCreatedAt}" in p.text:
                p.text = p.text.replace("${documentCreatedAt}", str(datetime.now().date()))

            if "${formattedChairmanOfCommision}" in p.text:
                p.text = p.text.replace(
                    "${formattedChairmanOfCommision}",
                    parsed_object['chairmanOfCommision']['name'].split()[1][0] + ". " +
                    parsed_object['chairmanOfCommision']['name'].split()[0]
                )

            if "${commissionMembersInstallationOrganization}" in p.text:
                resultString = ''
                for member in parsed_object['commissionMembers']:
                    resultString += (member['IIN'] + ', ' + member['name'] + ', ' + member['jobTitle'] + ', ' +
                                     member['BIN'] + ', ' + member['organizationName'] + '\n')
                p.text = p.text.replace("${commissionMembersInstallationOrganization}", resultString)

            if "${commissionMembersСommissioningOrganization}" in p.text:
                resultString = ''
                for member in parsed_object['commissionMembers']:
                    resultString += (member['IIN'] + ', ' + member['name'] + ', ' + member['jobTitle'] + ', ' +
                                     member['BIN'] + ', ' + member['organizationName'] + '\n')
                p.text = p.text.replace("${commissionMembersСommissioningOrganization}", resultString)

            if "${formattedCommisionMembers}" in p.text:
                resultString = ''
                for member in parsed_object['commissionMembers']:
                    resultString += member['name'].split()[1][0] + ". " + member['name'].split()[0] + '\n'
                p.text = p.text.replace("${formattedCommisionMembers}", resultString)

            document.save(file_path)

        return file_path
    except Exception as e:
        log_exception(e, 'Error in actOfCommissioningFireAutomationSystemsAndInstallations')
        return False


def reportOfTestingInternalFireFightingWaterSupply(data):
    try:
        json_string = json.dumps(data, default=convert_datetime, indent=2)
        parsed_object = json.loads(json_string)
        data_str = str(datetime.now().date())

        document = Document(
            f"{TMP_PATH}/Протокол испытаний внутреннего противопожарного водоснабжения на водоотдачу template.docx")
        file_path = f"{DOC_PATH}/Протокол испытаний внутреннего противопожарного водоснабжения на водоотдачу {data_str}.docx"

        document.styles.add_style("CommentsStyle", WD_STYLE_TYPE.PARAGRAPH)
        style = document.styles["CommentsStyle"]
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)

        for p in document.paragraphs:
            p.style = document.styles["CommentsStyle"]
            for key, value in parsed_object.items():
                placeholder = f"${{{key}}}"
                if type(value) is dict:
                    for key2, value2 in value.items():
                        placeholder = f"${{{key}.{key2}}}"
                        if placeholder in p.text:
                            p.text = p.text.replace(placeholder, str(value2))
                if "carriedOut" or "acceptedForOperationFrom" in key:
                    p.text = p.text.replace(
                        placeholder,
                        str(datetime.fromisoformat(value.replace("Z", "+00:00")).date())
                    )
                if placeholder in p.text:
                    p.text = p.text.replace(placeholder, str(value))

            if "${documentNumber}" in p.text:
                p.text = p.text.replace("${documentNumber}", str(uuid.uuid4()))

            if "${documentCreatedAt}" in p.text:
                p.text = p.text.replace("${documentCreatedAt}", str(datetime.now().date()))

            if "${formattedChairmanOfCommision}" in p.text:
                p.text = p.text.replace(
                    "${formattedChairmanOfCommision}",
                    parsed_object['chairmanOfCommision']['name'].split()[1][0] + ". " +
                    parsed_object['chairmanOfCommision']['name'].split()[0]
                )

            if "${commisionMembers}" in p.text:
                resultString = ''
                for member in parsed_object['commissionMembers']:
                    resultString += (member['IIN'] + ', ' + member['name'] + ', ' + member['jobTitle'] + ', ' +
                                     member['BIN'] + ', ' + member['organizationName'] + '\n')
                p.text = p.text.replace("${commisionMembers}", resultString)

            if "${formattedTestConductedMembers}" in p.text:
                resultString = ''
                for member in parsed_object['conductedMembers']:
                    resultString += member['name'].split()[1][0] + ". " + member['name'].split()[0] + '\n'
                p.text = p.text.replace("${formattedTestConductedMembers}", resultString)

            if "${testConductedMembers}" in p.text:
                resultString = ''
                for member in parsed_object['conductedMembers']:
                    resultString += member['name'].split()[1][0] + ". " + member['name'].split()[0] + '\n'
                p.text = p.text.replace("${testConductedMembers}", resultString)

            document.save(file_path)

        return file_path
    except Exception as e:
        log_exception(e, 'Error in reportOfTestingInternalFireFightingWaterSupply')
        return False


def statementOfInstalledDevicesAndEquipmentFireAutomationSystemsInstallations(data):
    try:
        json_string = json.dumps(data, default=convert_datetime, indent=2)
        parsed_object = json.loads(json_string)
        data_str = str(datetime.now().date())

        document = Document(
            f"{TMP_PATH}/ВЕДОМОСТЬ смонтированных приборов и оборудования систем и установок пожарной автоматики template.docx")
        file_path = f"{DOC_PATH}/ВЕДОМОСТЬ смонтированных приборов и оборудования систем и установок пожарной автоматики {data_str}.docx"

        document.styles.add_style("CommentsStyle", WD_STYLE_TYPE.PARAGRAPH)
        style = document.styles["CommentsStyle"]
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)

        for p in document.paragraphs:
            p.style = document.styles["CommentsStyle"]
            for key, value in parsed_object.items():
                placeholder = f"${{{key}}}"
                if type(value) is dict:
                    for key2, value2 in value.items():
                        placeholder = f"${{{key}.{key2}}}"
                        if placeholder in p.text:
                            p.text = p.text.replace(placeholder, str(value2))
                if "carriedOut" in key or "acceptedForOperationFrom" in key:
                    p.text = p.text.replace(
                        placeholder,
                        str(datetime.fromisoformat(value.replace("Z", "+00:00")).date())
                    )
                if placeholder in p.text:
                    p.text = p.text.replace(placeholder, str(value))

            if "${documentNumber}" in p.text:
                p.text = p.text.replace("${documentNumber}", str(uuid.uuid4()))

            if "${documentCreatedAt}" in p.text:
                p.text = p.text.replace("${documentCreatedAt}", str(datetime.now().date()))

            document.save(file_path)

        return file_path
    except Exception as e:
        log_exception(e, 'Error in statementOfInstalledDevicesAndEquipmentFireAutomationSystemsInstallations')
        return False


def testReportOfValvesFireCranesForOperability(data):
    try:
        json_string = json.dumps(data, default=convert_datetime, indent=2)
        parsed_object = json.loads(json_string)
        data_str = str(datetime.now().date())

        document = Document(
            f"{TMP_PATH}/Протокол испытаний клапанов пожарных кранов на работоспособность template.docx")
        file_path = f"{DOC_PATH}/Протокол испытаний клапанов пожарных кранов на работоспособность {data_str}.docx"

        document.styles.add_style("CommentsStyle", WD_STYLE_TYPE.PARAGRAPH)
        style = document.styles["CommentsStyle"]
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)

        for p in document.paragraphs:
            p.style = document.styles["CommentsStyle"]
            for key, value in parsed_object.items():
                placeholder = f"${{{key}}}"
                if type(value) is dict:
                    for key2, value2 in value.items():
                        placeholder = f"${{{key}.{key2}}}"
                        if placeholder in p.text:
                            p.text = p.text.replace(placeholder, str(value2))
                if "testTime" in key:
                    time_placeholder = "${testTime}"
                    if time_placeholder in p.text:
                        p.text = p.text.replace(
                            time_placeholder,
                            str(datetime.fromisoformat(value.replace("Z", "+00:00")).time())
                        )
                elif "testDate" in key:
                    date_placeholder = "${testDate}"
                    if date_placeholder in p.text:
                        p.text = p.text.replace(
                            date_placeholder,
                            str(datetime.fromisoformat(value.replace("Z", "+00:00")).date())
                        )

                if placeholder in p.text:
                    p.text = p.text.replace(placeholder, str(value))

            if "${documentNumber}" in p.text:
                p.text = p.text.replace("${documentNumber}", str(uuid.uuid4()))

            if "${documentCreatedAt}" in p.text:
                p.text = p.text.replace("${documentCreatedAt}", str(datetime.now().date()))

            if "${formattedChairmanOfCommision}" in p.text:
                p.text = p.text.replace(
                    "${formattedChairmanOfCommision}",
                    parsed_object['chairmanOfCommision']['name'].split()[1][0] + ". " +
                    parsed_object['chairmanOfCommision']['name'].split()[0]
                )

            if "${commisionMembers}" in p.text:
                resultString = ''
                for member in parsed_object['commissionMembers']:
                    resultString += (member['IIN'] + ', ' + member['name'] + ', ' + member['jobTitle'] + ', ' +
                                     member['BIN'] + ', ' + member['organizationName'] + '\n')
                p.text = p.text.replace("${commisionMembers}", resultString)

            if "${formattedTestConductedMembers}" in p.text:
                resultString = ''
                for member in parsed_object['conductedMembers']:
                    resultString += member['name'].split()[1][0] + ". " + member['name'].split()[0] + '\n'
                p.text = p.text.replace("${formattedTestConductedMembers}", resultString)

            if "${testConductedMembers}" in p.text:
                resultString = ''
                for member in parsed_object['conductedMembers']:
                    resultString += member['name'].split()[1][0] + ". " + member['name'].split()[0] + '\n'
                p.text = p.text.replace("${testConductedMembers}", resultString)

            document.save(file_path)

        return file_path
    except Exception as e:
        log_exception(e, 'Error in testReportOfValvesFireCranesForOperability')
        return False


def operationalLogOfFireAutomationSystemsInstallations(data):
    try:
        json_string = json.dumps(data, default=convert_datetime, indent=2)
        parsed_object = json.loads(json_string)
        data_str = str(datetime.now().date())

        document = Document(f"{TMP_PATH}/Эксплуатационный журнал систем и установок пожарной автоматики template.docx")
        file_path = f"{DOC_PATH}/Эксплуатационный журнал систем и установок пожарной автоматики {data_str}.docx"

        document.styles.add_style("CommentsStyle", WD_STYLE_TYPE.PARAGRAPH)
        style = document.styles["CommentsStyle"]
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)

        for p in document.paragraphs:
            p.style = document.styles["CommentsStyle"]
            for key, value in parsed_object.items():
                placeholder = f"${{{key}}}"
                if type(value) is dict:
                    for key2, value2 in value.items():
                        placeholder = f"${{{key}.{key2}}}"
                        if placeholder in p.text:
                            p.text = p.text.replace(placeholder, str(value2))
                if "carriedOut" in key or "acceptedForOperationFrom" in key:
                    p.text = p.text.replace(
                        placeholder,
                        str(datetime.fromisoformat(value.replace("Z", "+00:00")).date())
                    )
                if placeholder in p.text:
                    p.text = p.text.replace(placeholder, str(value))

            if "${documentNumber}" in p.text:
                p.text = p.text.replace("${documentNumber}", str(uuid.uuid4()))

            if "${documentCreatedAt}" in p.text:
                p.text = p.text.replace("${documentCreatedAt}", str(datetime.now().date()))

            document.save(file_path)

        return file_path
    except Exception as e:
        log_exception(e, 'Error in operationalLogOfFireAutomationSystemsInstallations')
        return False


def logbookOfTestingLaboratoryProtocols(data):
    try:
        json_string = json.dumps(data, default=convert_datetime, indent=2)
        parsed_object = json.loads(json_string)
        data_str = str(datetime.now().date())

        document = Document(f"{TMP_PATH}/Журнал регистрации протоколов испытательной лаборатории template.docx")
        file_path = f"{DOC_PATH}/Журнал регистрации протоколов испытательной лаборатории {data_str}.docx"

        document.styles.add_style("CommentsStyle", WD_STYLE_TYPE.PARAGRAPH)
        style = document.styles["CommentsStyle"]
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)

        for p in document.paragraphs:
            p.style = document.styles["CommentsStyle"]
            for key, value in parsed_object.items():
                placeholder = f"${{{key}}}"
                if type(value) is dict:
                    for key2, value2 in value.items():
                        placeholder = f"${{{key}.{key2}}}"
                        if placeholder in p.text:
                            p.text = p.text.replace(placeholder, str(value2))
                if "carriedOut" in key or "acceptedForOperationFrom" in key:
                    p.text = p.text.replace(
                        placeholder,
                        str(datetime.fromisoformat(value.replace("Z", "+00:00")).date())
                    )
                if placeholder in p.text:
                    p.text = p.text.replace(placeholder, str(value))

            if "${documentNumber}" in p.text:
                p.text = p.text.replace("${documentNumber}", str(uuid.uuid4()))

            if "${documentCreatedAt}" in p.text:
                p.text = p.text.replace("${documentCreatedAt}", str(datetime.now().date()))

            document.save(file_path)

        return file_path
    except Exception as e:
        log_exception(e, 'Error in logbookOfTestingLaboratoryProtocols')
        return False


def testReportTestingParametersOfVentilationSystems(data):
    try:
        json_string = json.dumps(data, default=convert_datetime, indent=2)
        parsed_object = json.loads(json_string)
        data_str = str(datetime.now().date())

        document = Document(f"{TMP_PATH}/Протокол испытаний испытание параметров систем вентиляции template.docx")
        file_path = f"{DOC_PATH}/Протокол испытаний испытание параметров систем вентиляции {data_str}.docx"

        document.styles.add_style("CommentsStyle", WD_STYLE_TYPE.PARAGRAPH)
        style = document.styles["CommentsStyle"]
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)

        for p in document.paragraphs:
            p.style = document.styles["CommentsStyle"]
            for key, value in parsed_object.items():
                placeholder = f"${{{key}}}"
                if type(value) is dict:
                    for key2, value2 in value.items():
                        placeholder = f"${{{key}.{key2}}}"
                        if placeholder in p.text:
                            p.text = p.text.replace(placeholder, str(value2))
                if "testDate" in key:
                    p.text = p.text.replace(
                        placeholder,
                        str(datetime.fromisoformat(value.replace("Z", "+00:00")).date())
                    )
                if placeholder in p.text:
                    p.text = p.text.replace(placeholder, str(value))

            if "${documentNumber}" in p.text:
                p.text = p.text.replace("${documentNumber}", str(uuid.uuid4()))

            if "${documentCreatedAt}" in p.text:
                p.text = p.text.replace("${documentCreatedAt}", str(datetime.now().date()))

            document.save(file_path)

        return file_path
    except Exception as e:
        log_exception(e, 'Error in testReportTestingParametersOfVentilationSystems')
        return False


def testsToDetermineTheStrengthOfFirefightingExternalStationaryLaddersAndRoofRailings(data):
    try:
        json_string = json.dumps(data, default=convert_datetime, indent=2)
        parsed_object = json.loads(json_string)
        data_str = str(datetime.now().date())

        document = Document(
            f"{TMP_PATH}/Испытания по определению прочностей пожарных наружных стационарных лестниц и ограждения кровли template.docx")
        file_path = f"{DOC_PATH}/Испытания по определению прочностей пожарных наружных стационарных лестниц и ограждения кровли {data_str}.docx"

        document.styles.add_style("CommentsStyle", WD_STYLE_TYPE.PARAGRAPH)
        style = document.styles["CommentsStyle"]
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)

        for p in document.paragraphs:
            p.style = document.styles["CommentsStyle"]
            for key, value in parsed_object.items():
                placeholder = f"${{{key}}}"
                if type(value) is dict:
                    for key2, value2 in value.items():
                        placeholder = f"${{{key}.{key2}}}"
                        if placeholder in p.text:
                            p.text = p.text.replace(placeholder, str(value2))
                if "testDate" in key:
                    p.text = p.text.replace(
                        placeholder,
                        str(datetime.fromisoformat(value.replace("Z", "+00:00")).date())
                    )
                if placeholder in p.text:
                    p.text = p.text.replace(placeholder, str(value))

            if "${documentNumber}" in p.text:
                p.text = p.text.replace("${documentNumber}", str(uuid.uuid4()))

            if "${documentCreatedAt}" in p.text:
                p.text = p.text.replace("${documentCreatedAt}", str(datetime.now().date()))

            document.save(file_path)

        return file_path
    except Exception as e:
        log_exception(e, 'Error in testsToDetermineTheStrengthOfFirefightingExternalStationaryLaddersAndRoofRailings')
        return False


def protocolTestLightningProtectionTestingSystem(data):
    try:
        json_string = json.dumps(data, default=convert_datetime, indent=2)
        parsed_object = json.loads(json_string)
        data_str = str(datetime.now().date())

        document = Document(f"{TMP_PATH}/ПРОТОКОЛ ИСПЫТАНИЙ проверки систем молниезащиты template.docx")
        file_path = f"{DOC_PATH}/ПРОТОКОЛ ИСПЫТАНИЙ проверки систем молниезащиты {data_str}.docx"

        document.styles.add_style("CommentsStyle", WD_STYLE_TYPE.PARAGRAPH)
        style = document.styles["CommentsStyle"]
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)

        for p in document.paragraphs:
            p.style = document.styles["CommentsStyle"]
            for key, value in parsed_object.items():
                placeholder = f"${{{key}}}"
                if type(value) is dict:
                    for key2, value2 in value.items():
                        placeholder = f"${{{key}.{key2}}}"
                        if placeholder in p.text:
                            p.text = p.text.replace(placeholder, str(value2))
                if "testDate" in key:
                    p.text = p.text.replace(
                        placeholder,
                        str(datetime.fromisoformat(value.replace("Z", "+00:00")).date())
                    )
                if placeholder in p.text:
                    p.text = p.text.replace(placeholder, str(value))

            if "${documentNumber}" in p.text:
                p.text = p.text.replace("${documentNumber}", str(uuid.uuid4()))

            if "${documentCreatedAt}" in p.text:
                p.text = p.text.replace("${documentCreatedAt}", str(datetime.now().date()))

            document.save(file_path)

        return file_path
    except Exception as e:
        log_exception(e, 'Error in protocolTestLightningProtectionTestingSystem')
        return False


def formOfControlTestReportToDetermineTheQualityOfFireRetardantForMetalStructures(data):
    try:
        json_string = json.dumps(data, default=convert_datetime, indent=2)
        parsed_object = json.loads(json_string)
        data_str = str(datetime.now().date())

        document = Document(
            f"{TMP_PATH}/Форма протокола контрольных испытаний по определению качества обработки объекта огнезащитной по металлоконструкциям template.docx")
        file_path = f"{DOC_PATH}/Форма протокола контрольных испытаний по определению качества обработки объекта огнезащитной по металлоконструкциям {data_str}.docx"

        document.styles.add_style("CommentsStyle", WD_STYLE_TYPE.PARAGRAPH)
        style = document.styles["CommentsStyle"]
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)

        for p in document.paragraphs:
            p.style = document.styles["CommentsStyle"]
            for key, value in parsed_object.items():
                placeholder = f"${{{key}}}"
                if type(value) is dict:
                    for key2, value2 in value.items():
                        placeholder = f"${{{key}.{key2}}}"
                        if placeholder in p.text:
                            p.text = p.text.replace(placeholder, str(value2))
                if "Date" in key:
                    p.text = p.text.replace(placeholder,
                                            str(datetime.fromisoformat(value.replace("Z", "+00:00")).date()))
                if placeholder in p.text:
                    p.text = p.text.replace(placeholder, str(value))

            if "${documentNumber}" in p.text:
                p.text = p.text.replace("${documentNumber}", str(uuid.uuid4()))

            if "${documentCreatedAt}" in p.text:
                p.text = p.text.replace("${documentCreatedAt}", str(datetime.now().date()))

            if "${formattedTestConducted}" in p.text:
                p.text = p.text.replace(
                    "${formattedTestConducted}",
                    parsed_object['TestConducted']['name'].split()[1][0] + ". " +
                    parsed_object['TestConducted']['name'].split()[0]
                )

            document.save(file_path)

        return file_path
    except Exception as e:
        log_exception(e, 'Error in formOfControlTestReportToDetermineTheQualityOfFireRetardantForMetalStructures')
        return False


def formOfControlTestReportToDetermineTheQualityOfFireRetardantForWoodenStructures(data):
    try:
        json_string = json.dumps(data, default=convert_datetime, indent=2)
        parsed_object = json.loads(json_string)
        data_str = str(datetime.now().date())

        document = Document(
            f"{TMP_PATH}/Форма протокола контрольных испытаний по определению качества обработки объекта огнезащитной по деревянным конструкциям template.docx")
        file_path = f"{DOC_PATH}/Форма протокола контрольных испытаний по определению качества обработки объекта огнезащитной по деревянным конструкциям {data_str}.docx"

        document.styles.add_style("CommentsStyle", WD_STYLE_TYPE.PARAGRAPH)
        style = document.styles["CommentsStyle"]
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)

        for p in document.paragraphs:
            p.style = document.styles["CommentsStyle"]
            for key, value in parsed_object.items():
                placeholder = f"${{{key}}}"
                if type(value) is dict:
                    for key2, value2 in value.items():
                        placeholder = f"${{{key}.{key2}}}"
                        if placeholder in p.text:
                            p.text = p.text.replace(placeholder, str(value2))
                if "Date" in key:
                    p.text = p.text.replace(placeholder,
                                            str(datetime.fromisoformat(value.replace("Z", "+00:00")).date()))
                if placeholder in p.text:
                    p.text = p.text.replace(placeholder, str(value))

            if "${documentNumber}" in p.text:
                p.text = p.text.replace("${documentNumber}", str(uuid.uuid4()))

            if "${documentCreatedAt}" in p.text:
                p.text = p.text.replace("${documentCreatedAt}", str(datetime.now().date()))

            if "${formattedTestConducted}" in p.text:
                p.text = p.text.replace(
                    "${formattedTestConducted}",
                    parsed_object['TestConducted']['name'].split()[1][0] + ". " +
                    parsed_object['TestConducted']['name'].split()[0]
                )

            document.save(file_path)

        return file_path
    except Exception as e:
        log_exception(e, 'Error in formOfControlTestReportToDetermineTheQualityOfFireRetardantForWoodenStructures')
        return False


def TESTREPORTFlammabilityTestMethodAndClassification(data):
    try:
        json_string = json.dumps(data, default=convert_datetime, indent=2)
        parsed_object = json.loads(json_string)
        data_str = str(datetime.now().date())

        document = Document(
            f"{TMP_PATH}/ПРОТОКОЛ ИСПЫТАНИЙ метод испытания на воспламеняемость и классификация template.docx")
        file_path = f"{DOC_PATH}/ПРОТОКОЛ ИСПЫТАНИЙ метод испытания на воспламеняемость и классификация {data_str}.docx"

        document.styles.add_style("CommentsStyle", WD_STYLE_TYPE.PARAGRAPH)
        style = document.styles["CommentsStyle"]
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)

        for p in document.paragraphs:
            p.style = document.styles["CommentsStyle"]
            for key, value in parsed_object.items():
                placeholder = f"${{{key}}}"
                if type(value) is dict:
                    for key2, value2 in value.items():
                        placeholder = f"${{{key}.{key2}}}"
                        if placeholder in p.text:
                            p.text = p.text.replace(placeholder, str(value2))
                if "Date" in key:
                    p.text = p.text.replace(placeholder,
                                            str(datetime.fromisoformat(value.replace("Z", "+00:00")).date()))
                if placeholder in p.text:
                    p.text = p.text.replace(placeholder, str(value))

            if "${documentNumber}" in p.text:
                p.text = p.text.replace("${documentNumber}", str(uuid.uuid4()))

            if "${documentCreatedAt}" in p.text:
                p.text = p.text.replace("${documentCreatedAt}", str(datetime.now().date()))

            if "${formattedTestConducted}" in p.text:
                p.text = p.text.replace(
                    "${formattedTestConducted}",
                    parsed_object['TestConducted']['name'].split()[1][0] + ". " +
                    parsed_object['TestConducted']['name'].split()[0]
                )

            document.save(file_path)

        return file_path
    except Exception as e:
        log_exception(e, 'Error in TESTREPORTFlammabilityTestMethodAndClassification')
        return False


def TESTREPORTForMeasuringTheInsulationResistanceOfWiresAndCables(data):
    try:
        json_string = json.dumps(data, default=convert_datetime, indent=2)
        parsed_object = json.loads(json_string)
        data_str = str(datetime.now().date())

        document = Document(
            f"{TMP_PATH}/ПРОТОКОЛ ИСПЫТАНИЙ измерения сопротивления изоляции проводов и кабелей template.docx")
        file_path = f"{DOC_PATH}/ПРОТОКОЛ ИСПЫТАНИЙ измерения сопротивления изоляции проводов и кабелей {data_str}.docx"

        document.styles.add_style("CommentsStyle", WD_STYLE_TYPE.PARAGRAPH)
        style = document.styles["CommentsStyle"]
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)

        for p in document.paragraphs:
            p.style = document.styles["CommentsStyle"]
            for key, value in parsed_object.items():
                placeholder = f"${{{key}}}"
                if type(value) is dict:
                    for key2, value2 in value.items():
                        placeholder = f"${{{key}.{key2}}}"
                        if placeholder in p.text:
                            p.text = p.text.replace(placeholder, str(value2))
                if "Date" in key:
                    p.text = p.text.replace(placeholder,
                                            str(datetime.fromisoformat(value.replace("Z", "+00:00")).date()))
                if placeholder in p.text:
                    p.text = p.text.replace(placeholder, str(value))

            if "${documentNumber}" in p.text:
                p.text = p.text.replace("${documentNumber}", str(uuid.uuid4()))

            if "${documentCreatedAt}" in p.text:
                p.text = p.text.replace("${documentCreatedAt}", str(datetime.now().date()))

            if "${formattedTestConducted}" in p.text:
                p.text = p.text.replace(
                    "${formattedTestConducted}",
                    parsed_object['TestConducted']['name'].split()[1][0] + ". " +
                    parsed_object['TestConducted']['name'].split()[0]
                )

            document.save(file_path)

        return file_path
    except Exception as e:
        log_exception(e, 'Error in TESTREPORTForMeasuringTheInsulationResistanceOfWiresAndCables')
        return False
